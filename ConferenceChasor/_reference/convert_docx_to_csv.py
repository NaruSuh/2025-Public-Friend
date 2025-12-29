#!/usr/bin/env python3
"""
워드 파일의 표를 CSV로 변환하는 유틸리티

사용법:
    python convert_docx_to_csv.py input.docx output.csv
    python convert_docx_to_csv.py input.docx output.csv --table 0
"""

import argparse
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / 'src'))

from docx_processor import DocxDataReader


def main():
    parser = argparse.ArgumentParser(
        description='워드 파일의 표를 CSV로 변환',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
예시:
  # 첫 번째 표를 CSV로 변환
  python convert_docx_to_csv.py data/survey.docx data/survey.csv

  # 두 번째 표를 CSV로 변환 (인덱스 1)
  python convert_docx_to_csv.py data/survey.docx data/survey.csv --table 1

  # 모든 표 확인
  python convert_docx_to_csv.py data/survey.docx --list
        """
    )

    parser.add_argument('input', help='입력 워드 파일 (.docx)')
    parser.add_argument('output', nargs='?', help='출력 CSV 파일')
    parser.add_argument('--table', type=int, default=0,
                       help='변환할 표 인덱스 (기본값: 0, 첫 번째 표)')
    parser.add_argument('--list', action='store_true',
                       help='워드 파일의 모든 표 목록 보기')

    args = parser.parse_args()

    # 입력 파일 확인
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"오류: 입력 파일을 찾을 수 없습니다: {args.input}")
        sys.exit(1)

    try:
        reader = DocxDataReader(str(input_path))

        if args.list:
            # 모든 표 목록 출력
            print(f"\n'{args.input}' 파일의 표 목록:")
            print("=" * 60)

            from docx import Document
            doc = Document(str(input_path))
            print(f"총 표 개수: {len(doc.tables)}개\n")

            for i, table in enumerate(doc.tables):
                print(f"표 {i}:")
                print(f"  - 행 수: {len(table.rows)}")
                print(f"  - 열 수: {len(table.columns)}")
                if len(table.rows) > 0:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]
                    print(f"  - 헤더: {headers[:5]}...")  # 처음 5개만
                print()

            print("=" * 60)
            print("\n변환하려면:")
            print(f"  python convert_docx_to_csv.py {args.input} output.csv --table <인덱스>")
        else:
            # CSV로 변환
            if not args.output:
                print("오류: 출력 파일 경로를 지정해주세요")
                sys.exit(1)

            print(f"워드 파일 읽기: {args.input}")
            print(f"표 인덱스: {args.table}")

            df = reader.extract_table_data(args.table)

            print(f"\n추출된 데이터:")
            print(f"  - 행 수: {len(df)}")
            print(f"  - 열 수: {len(df.columns)}")
            print(f"  - 컬럼: {list(df.columns)}")

            # CSV 저장
            df.to_csv(args.output, index=False, encoding='utf-8-sig')
            print(f"\nCSV 저장 완료: {args.output}")

            print("\n이제 다음 명령으로 보고서를 생성할 수 있습니다:")
            print(f"  python generate_report.py --data {args.output} --config config.json")

    except Exception as e:
        print(f"\n오류 발생: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
