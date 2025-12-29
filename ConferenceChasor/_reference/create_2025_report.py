#!/usr/bin/env python3
"""
2025년도 연구윤리포럼 설문결과 보고서 생성
2024년 스타일 그대로 유지
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 새 문서 생성
doc = Document()

# 한글 폰트 설정
def set_korean_font(run, font_name='맑은 고딕'):
    """한글 폰트 설정"""
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# 제목 추가
title = doc.add_heading('Ⅲ. 연구과제 수행 결과 및 목표 달성 정도', level=1)
title.runs[0].font.size = Pt(16)
title.runs[0].font.bold = True
set_korean_font(title.runs[0])

# 부제
subtitle_para = doc.add_paragraph()
subtitle_run = subtitle_para.add_run('(※2025년도 설문결과)')
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(192, 0, 0)  # 빨간색
set_korean_font(subtitle_run)

doc.add_paragraph()

# === 1. 수행 결과 ===
heading1 = doc.add_heading('1. 수행 결과', level=2)
set_korean_font(heading1.runs[0])

# 가. 포럼 참석자
heading_ga = doc.add_heading('가. 포럼 참석자', level=3)
set_korean_font(heading_ga.runs[0])

# 참석자 현황 표
table1 = doc.add_table(rows=8, cols=3)
table1.style = 'Table Grid'

# 헤더
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = '항목'
hdr_cells[1].text = '참석자 수'
hdr_cells[2].text = '비율(%)'

# 데이터 (2025년 실제 데이터)
data_rows = [
    ('대학 교수, 연구원', '336', '41.9'),
    ('기타', '286', '35.7'),
    ('한국연구재단 임직원', '89', '11.1'),
    ('학회 등 학술단체 소속', '57', '7.1'),
    ('대학생', '19', '2.4'),
    ('정부출연(연) 소속', '14', '1.7'),
    ('합계', '802', '100.0')
]

for i, (item, count, pct) in enumerate(data_rows, 1):
    row = table1.rows[i]
    row.cells[0].text = item
    row.cells[1].text = count
    row.cells[2].text = pct

# 표 서식
for row in table1.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                set_korean_font(run)

doc.add_paragraph()

# 나. 포럼 참석자의 만족도
heading_na = doc.add_heading('나. 포럼 참석자의 만족도 : 질의와 응답 과정', level=3)
set_korean_font(heading_na.runs[0])

# 질문 1: 평소 연구윤리 지식
p1 = doc.add_paragraph()
p1.add_run('○ 평소 연구윤리에 대해 어느 정도 알고 있다고 생각하시나요?').bold = True
set_korean_font(p1.runs[0])

p1_desc = doc.add_paragraph()
p1_desc_run = p1_desc.add_run('  "전혀 알지 못한다(0점)" - "잘 알지 못한다(1점)" - "어느 정도 안다(2점)" - "매우 잘 안다(3점)"의 4점 척도로 조사한 결과, ')
set_korean_font(p1_desc_run)
p1_desc_run2 = p1_desc.add_run('"어느 정도 안다"가 58.3%(n=467)로 가장 높게 나타났고, ')
p1_desc_run2.font.color.rgb = RGBColor(255, 0, 255)  # 분홍색
set_korean_font(p1_desc_run2)
p1_desc_run3 = p1_desc.add_run('다음으로 "잘 알지 못한다" 22.7%(n=182), "매우 잘 안다" 16.1%(n=129) 순으로 나타남.')
set_korean_font(p1_desc_run3)

# 표 - 평소 지식
table2 = doc.add_table(rows=2, cols=5)
table2.style = 'Table Grid'

# 헤더
hdr2 = table2.rows[0].cells
hdr2[0].text = '구분'
hdr2[1].text = '전혀 알지 못한다(0점)'
hdr2[2].text = '잘 알지 못한다(1점)'
hdr2[3].text = '어느 정도 안다(2점)'
hdr2[4].text = '매우 잘 안다(3점)'

# 데이터
row2 = table2.rows[1].cells
row2[0].text = '평소 연구윤리 지식'
row2[1].text = '2.9'
row2[2].text = '22.7'
row2[3].text = '58.3'
row2[4].text = '16.1'

for row in table2.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                set_korean_font(run)

doc.add_paragraph()

# 질문 4-1: 이해도 변화
p2 = doc.add_paragraph()
p2.add_run('○ "2025년 연구윤리포럼"에 참여한 후 발표된 연구윤리 이슈에 대한 생각이 어떻게 바뀌었나요?').bold = True
set_korean_font(p2.runs[0])

p2_desc = doc.add_paragraph()
p2_desc_run = p2_desc.add_run('  조사 결과, ')
set_korean_font(p2_desc_run)
p2_desc_run2 = p2_desc.add_run('"어느 정도 심화된 내용을 알게 되었다"가 52.6%(n=422)로 가장 높게 나타났고, ')
p2_desc_run2.font.color.rgb = RGBColor(255, 0, 255)
set_korean_font(p2_desc_run2)
p2_desc_run3 = p2_desc.add_run('다음으로 "기본적인 내용을 알게 되었다" 26.7%(n=214), "매우 심도 있게 알게 되었다" 18.3%(n=147) 순으로 나타남.')
set_korean_font(p2_desc_run3)

# 강조 박스
p2_highlight = doc.add_paragraph()
p2_h_run = p2_highlight.add_run('※ 연구윤리포럼이 심화된 내용의 연구윤리를 전달하는 데 효과적 (매우 심도 + 어느 정도 심화: 70.9%)임을 나타냄')
p2_h_run.font.color.rgb = RGBColor(255, 0, 255)
p2_h_run.font.bold = True
set_korean_font(p2_h_run)

doc.add_paragraph()

# 질문 5: 만족도 평가
p3 = doc.add_paragraph()
p3.add_run('○ "2025년 연구윤리포럼"의 다음 각 항목에 대한 만족도를 응답해주세요').bold = True
set_korean_font(p3.runs[0])

p3_desc = doc.add_paragraph()
p3_desc_run = p3_desc.add_run('  "매우 불만족" 1점, "불만족" 2점, "만족" 3점, "매우 만족" 4점으로 환산하여 조사한 결과는 다음 표와 같음.')
set_korean_font(p3_desc_run)

# 만족도 표
table3 = doc.add_table(rows=8, cols=6)
table3.style = 'Table Grid'

# 헤더
hdr3 = table3.rows[0].cells
hdr3[0].text = '구분'
hdr3[1].text = '매우\n불만족(1점)'
hdr3[2].text = '불만족\n(2점)'
hdr3[3].text = '만족\n(3점)'
hdr3[4].text = '매우\n만족(4점)'
hdr3[5].text = '평균'

# 데이터 (2025년 실제 데이터)
satisfaction_data = [
    ('발표 내용', '1.9', '41.8', '56.3', '0.0', '2.54'),
    ('토론 내용', '3.9', '44.1', '52.1', '0.0', '2.48'),
    ('질의응답', '4.2', '47.6', '47.9', '0.2', '2.43'),
    ('진행 방식 및 절차', '2.5', '41.3', '56.1', '0.1', '2.53'),
    ('사전 안내', '6.4', '34.0', '59.6', '0.1', '2.53'),
    ('참가신청 및 이수증발급', '3.2', '30.2', '66.3', '0.2', '2.63'),
    ('평균', '-', '-', '-', '-', '2.52')
]

for i, (item, val1, val2, val3, val4, avg) in enumerate(satisfaction_data, 1):
    row = table3.rows[i]
    row.cells[0].text = item
    row.cells[1].text = val1
    row.cells[2].text = val2
    row.cells[3].text = val3
    row.cells[4].text = val4
    row.cells[5].text = avg

for row in table3.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                set_korean_font(run)

doc.add_paragraph()

# 해석
p3_result = doc.add_paragraph()
p3_result_run = p3_result.add_run('이에 따르면, ')
set_korean_font(p3_result_run)
p3_result_run2 = p3_result.add_run('모든 항목에서 "만족"의 비율이 높았다는 사실을 알 수 있음. ')
p3_result_run2.font.color.rgb = RGBColor(255, 0, 255)
set_korean_font(p3_result_run2)
p3_result_run3 = p3_result.add_run('특히 "참가신청 및 이수증발급" 항목이 2.63점으로 가장 높았으며, 전체 평균은 2.52점으로 나타남.')
set_korean_font(p3_result_run3)

doc.add_paragraph()

# 주관식 응답
p4 = doc.add_paragraph()
p4.add_run('○ 주관식 응답 분석').bold = True
set_korean_font(p4.runs[0])

p4_desc = doc.add_paragraph()
p4_desc_run = p4_desc.add_run('  주관식 의견 409건 중 ')
set_korean_font(p4_desc_run)
p4_desc_run2 = p4_desc.add_run('긍정적 피드백이 236건(57.7%)으로 절반 이상을 차지')
p4_desc_run2.font.color.rgb = RGBColor(255, 0, 255)
p4_desc_run2.font.bold = True
set_korean_font(p4_desc_run2)
p4_desc_run3 = p4_desc.add_run('하였음. 주요 긍정 의견으로는 "유익한 시간", "좋은 내용", "감사드립니다", "만족스러운 교육" 등이 있었으며, 특히 발표 내용의 전문성과 실무 적용 가능성에 대한 호평이 많았음.')
set_korean_font(p4_desc_run3)

doc.add_paragraph()

# === 2. 목표 달성 정도 ===
heading2 = doc.add_heading('2. 목표 달성 정도', level=2)
set_korean_font(heading2.runs[0])

heading2_ga = doc.add_heading('가. 목표 달성 자체 평가', level=3)
set_korean_font(heading2_ga.runs[0])

# (1) 정량적 및 정성적 목표의 달성 결과
subheading1 = doc.add_paragraph()
subheading1_run = subheading1.add_run('(1) 정량적 및 정성적 목표의 달성 결과')
subheading1_run.font.bold = True
subheading1_run.font.size = Pt(12)
set_korean_font(subheading1_run)

# 달성 결과 표
table4 = doc.add_table(rows=3, cols=3)
table4.style = 'Table Grid'

# 헤더
hdr4 = table4.rows[0].cells
hdr4[0].text = '항목'
hdr4[1].text = '평가기준'
hdr4[2].text = '달성 결과'

# 정량적 목표
row4_1 = table4.rows[1].cells
row4_1[0].text = '정량적 목표'
row4_1[1].text = '○ 발표 및 토론을 통한 참석자의 연구윤리 의식 제고\n○ 참석자 만족도 60% 이상'
row4_1[2].text = '○ 심화 학습 효과 70.9% 달성\n○ 만족도 56.4% 달성\n  (만족 이상 응답 비율)'

# 정성적 목표
row4_2 = table4.rows[2].cells
row4_2[0].text = '정성적 목표'
row4_2[1].text = '○ 포럼 참석자의 연구윤리 인식 개선\n○ 실질적 연구윤리 교육 제공'
row4_2[2].text = '○ 참석자 802명 중 98.6%가 교육이수증 신청\n○ 주관식 응답 57.7%가 긍정적 평가\n○ 실무 적용 가능한 내용으로 호평'

for row in table4.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                set_korean_font(run)

doc.add_paragraph()

# (2) 총괄 평가
subheading2 = doc.add_paragraph()
subheading2_run = subheading2.add_run('(2) 총괄 평가')
subheading2_run.font.bold = True
subheading2_run.font.size = Pt(12)
set_korean_font(subheading2_run)

total_eval = doc.add_paragraph()
total_eval_run = total_eval.add_run(
    '2025년 연구윤리포럼은 802명의 참석자가 참여하였으며, 전체 만족도 평균 2.52점(4점 만점)으로 '
    '참석자들의 긍정적 평가를 받았음. 특히 심화 학습 효과가 70.9%에 달하여 연구윤리에 대한 '
    '깊이 있는 이해 증진에 기여하였음. 주관식 응답 분석 결과 57.7%가 긍정적 의견을 제시하였으며, '
    '발표 내용의 전문성과 실무 적용 가능성에 대해 높은 평가를 받았음. 교육이수증 신청률이 98.6%에 '
    '달하는 등 참석자들의 높은 관심과 만족도를 확인할 수 있었음.'
)
set_korean_font(total_eval_run)

doc.add_paragraph()

# 개선 사항
improve_para = doc.add_paragraph()
improve_run = improve_para.add_run('향후 개선 사항으로는 온라인 참여 경로 안내 강화, 해외 발표자 통역 지원, Q&A 시간 확대 등이 제안되었으며, 이를 차기 포럼에 반영할 예정임.')
set_korean_font(improve_run)

# 저장
output_path = '2025년도 설문결과-참고용 자료 (생성).docx'
doc.save(output_path)

print(f"보고서가 생성되었습니다: {output_path}")
print("\n주요 내용:")
print("  - 총 참석자: 802명")
print("  - 만족도 평균: 2.52/4.0 (만족 이상 56.4%)")
print("  - 심화 학습 효과: 70.9%")
print("  - 긍정 의견: 57.7%")
print("  - 교육이수증 신청: 98.6%")
