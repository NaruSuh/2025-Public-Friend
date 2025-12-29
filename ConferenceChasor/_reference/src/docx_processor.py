"""
워드 파일 처리 모듈
- 워드 파일에서 설문 데이터 읽기
- 보고서를 워드 파일로 출력
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from typing import Dict, List, Any, Optional
from pathlib import Path
import re


class DocxDataReader:
    """워드 파일에서 설문조사 데이터 읽기"""

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = Document(str(self.docx_path))

    def extract_table_data(self, table_index: int = 0) -> pd.DataFrame:
        """워드 파일의 표에서 데이터 추출"""
        if table_index >= len(self.doc.tables):
            raise ValueError(f"표 인덱스 {table_index}가 범위를 벗어났습니다")

        table = self.doc.tables[table_index]
        data = []

        # 첫 행을 헤더로 사용
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # 나머지 행을 데이터로 수집
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):  # 빈 행 건너뛰기
                data.append(row_data)

        df = pd.DataFrame(data, columns=headers)
        return df

    def extract_all_tables(self) -> List[pd.DataFrame]:
        """워드 파일의 모든 표 추출"""
        tables = []
        for i in range(len(self.doc.tables)):
            try:
                df = self.extract_table_data(i)
                tables.append(df)
            except Exception as e:
                print(f"표 {i} 추출 중 오류: {e}")
        return tables

    def extract_structured_responses(self) -> pd.DataFrame:
        """구조화된 설문 응답 추출 (질문-응답 패턴)"""
        responses = []
        current_response = {}

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()

            if not text:
                # 빈 줄은 응답 구분자로 사용
                if current_response:
                    responses.append(current_response.copy())
                    current_response = {}
                continue

            # "질문: 답변" 또는 "Q1. 답변" 패턴 매칭
            match = re.match(r'(.+?)[:\.](.+)', text)
            if match:
                question = match.group(1).strip()
                answer = match.group(2).strip()
                current_response[question] = answer

        # 마지막 응답 추가
        if current_response:
            responses.append(current_response)

        return pd.DataFrame(responses)

    def save_as_csv(self, output_path: str, table_index: int = 0):
        """워드 표를 CSV로 저장"""
        df = self.extract_table_data(table_index)
        df.to_csv(output_path, index=False, encoding='utf-8-sig')
        print(f"CSV 저장 완료: {output_path}")


class DocxReportWriter:
    """보고서를 워드 파일로 작성"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """문서 스타일 설정"""
        # 기본 폰트 설정
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)

    def add_title(self, title: str, subtitle: str = ""):
        """제목 추가"""
        # 메인 제목
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 부제목
        if subtitle:
            p = self.doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    def add_metadata(self, metadata: Dict[str, str]):
        """메타데이터 추가 (조사 기간, 응답자 수 등)"""
        table = self.doc.add_table(rows=len(metadata), cols=2)
        table.style = 'Light Grid Accent 1'

        for i, (key, value) in enumerate(metadata.items()):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = str(value)

            # 키 셀 볼드 처리
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        self.doc.add_paragraph()  # 여백

    def add_section(self, title: str, content: str, level: int = 1):
        """섹션 추가"""
        self.doc.add_heading(title, level=level)

        if content:
            p = self.doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_highlighted_box(self, content: str, bg_color: tuple = (236, 240, 241)):
        """강조 박스 추가 (경영진 요약 등)"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 배경색은 직접 적용이 어려우므로 테이블로 구현
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = content

        # 셀 여백 설정
        cell_para = cell.paragraphs[0]
        cell_para.paragraph_format.space_before = Pt(6)
        cell_para.paragraph_format.space_after = Pt(6)

        self.doc.add_paragraph()  # 여백

    def add_stats_table(self, title: str, stats: Dict[str, Any]):
        """통계 표 추가"""
        self.doc.add_heading(title, level=2)

        # 표 생성
        table = self.doc.add_table(rows=len(stats) + 1, cols=2)
        table.style = 'Light List Accent 1'

        # 헤더
        header_cells = table.rows[0].cells
        header_cells[0].text = '항목'
        header_cells[1].text = '값'

        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # 데이터
        for i, (key, value) in enumerate(stats.items(), 1):
            row = table.rows[i]
            row.cells[0].text = str(key)

            # 값 포맷팅
            if isinstance(value, float):
                row.cells[1].text = f"{value:.2f}"
            else:
                row.cells[1].text = str(value)

        self.doc.add_paragraph()  # 여백

    def add_satisfaction_item(self, item_name: str, stats: Dict[str, Any],
                             analysis_text: str = ""):
        """만족도 항목 추가"""
        self.doc.add_heading(item_name, level=3)

        # 주요 지표를 표로 추가
        key_stats = {
            '평균 점수': f"{stats.get('mean', 0):.2f}/5.0",
            '중앙값': f"{stats.get('median', 0):.2f}",
            '표준편차': f"{stats.get('std', 0):.2f}",
            '응답 수': f"{stats.get('count', 0)}명"
        }

        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Light Grid Accent 1'

        # 첫 행: 항목명
        for i, key in enumerate(key_stats.keys()):
            cell = table.rows[0].cells[i]
            cell.text = key
            cell.paragraphs[0].runs[0].font.bold = True

        # 둘째 행: 값
        for i, value in enumerate(key_stats.values()):
            table.rows[1].cells[i].text = value

        # 분포 정보
        if 'distribution_percent' in stats:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph('응답 분포:')
            p.runs[0].font.bold = True

            for score, percent in sorted(stats['distribution_percent'].items()):
                bar = '█' * int(percent / 5)
                self.doc.add_paragraph(f"{score}점: {percent:.1f}% {bar}", style='List Bullet')

        # 분석 텍스트
        if analysis_text:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run('분석: ')
            run.font.bold = True
            p.add_run(analysis_text)

        self.doc.add_paragraph()  # 여백

    def add_text_responses(self, title: str, responses: List[str],
                          insight: str = "", max_display: int = 5):
        """주관식 응답 추가"""
        self.doc.add_heading(title, level=3)

        p = self.doc.add_paragraph()
        run = p.add_run(f'응답 수: ')
        run.font.bold = True
        p.add_run(f'{len(responses)}개')

        # 인사이트
        if insight:
            self.doc.add_paragraph()
            self.add_highlighted_box(insight)

        # 주요 응답 샘플
        if responses:
            self.doc.add_paragraph('주요 응답 샘플:', style='List Bullet')
            for i, resp in enumerate(responses[:max_display], 1):
                self.doc.add_paragraph(f'"{resp}"', style='List Bullet 2')

        self.doc.add_paragraph()  # 여백

    def add_recommendations(self, recommendations: str):
        """권장사항 추가"""
        self.doc.add_heading('개선 권장사항', level=2)

        # 권장사항이 불릿 포인트 형식인지 확인
        lines = recommendations.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                # '-' 또는 '•'로 시작하면 불릿 포인트로 추가
                if line.startswith(('-', '•', '*')):
                    self.doc.add_paragraph(line[1:].strip(), style='List Bullet')
                else:
                    self.doc.add_paragraph(line)

        self.doc.add_paragraph()  # 여백

    def add_page_break(self):
        """페이지 나누기"""
        self.doc.add_page_break()

    def save(self, output_path: str):
        """문서 저장"""
        self.doc.save(output_path)
        print(f"워드 문서 저장 완료: {output_path}")


def convert_docx_to_csv(docx_path: str, csv_path: str, table_index: int = 0):
    """워드 표를 CSV로 변환하는 유틸리티 함수"""
    reader = DocxDataReader(docx_path)
    reader.save_as_csv(csv_path, table_index)


def create_sample_docx_report(output_path: str):
    """샘플 워드 보고서 생성"""
    writer = DocxReportWriter()

    writer.add_title("학회 참석자 만족도 조사 보고서", "2024년 AI 국제학회")

    metadata = {
        '조사 기간': '2024년 11월',
        '응답자 수': '20명',
        '보고서 생성일': '2024년 11월 17일'
    }
    writer.add_metadata(metadata)

    writer.add_section('경영진 요약',
                      '본 조사는 2024년 AI 국제학회 참석자들의 만족도를 측정하였으며, '
                      '전체 만족도 지수는 85.5/100으로 매우 우수한 수준을 기록했습니다.')

    writer.add_stats_table('주요 지표', {
        '전체 만족도 지수': '85.5/100',
        '평가': '매우 우수'
    })

    writer.save(output_path)
